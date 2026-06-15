"""Debug: 检查合卷文件中答案的存储格式"""
import sys, re
sys.stdout.reconfigure(encoding="utf-8", errors="replace")
from docx import Document

# 检查一个合卷文件
path = "exam-papers/5.深圳中考化学（2008-2025）少22/2008年广东省深圳市中考化学真题及答案.docx"
print(f"文件: {path}")
doc = Document(path)

# 找包含"答案"关键词的段落
print("\n=== 含'答案'的段落 ===")
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if "答案" in text and len(text) < 100:
        print(f"  [{i}] {text}")

# 找题目后的内容
print("\n=== 前3道题及其后续内容 ===")
q_count = 0
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if re.match(r'^\d+[\.\、．]', text) and len(text) > 10:
        if q_count < 3:
            print(f"\n  --- 题 {q_count+1} (段落{i}) ---")
            print(f"  Q: {text[:200]}")
            # 显示后续5个段落
            for j in range(i+1, min(i+6, len(doc.paragraphs))):
                t2 = doc.paragraphs[j].text.strip()
                if t2:
                    print(f"    后续[{j}]: {t2[:150]}")
        q_count += 1

# 也检查表格
print(f"\n=== 表格数: {len(doc.tables)} ===")
for ti, table in enumerate(doc.tables[:2]):
    print(f"\n  表格 {ti}: {len(table.rows)}行 × {len(table.columns)}列")
    for ri, row in enumerate(table.rows[:5]):
        cells = [c.text.strip()[:50] for c in row.cells]
        print(f"    行{ri}: {cells}")
