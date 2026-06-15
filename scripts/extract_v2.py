# -*- coding: utf-8 -*-
"""
中考真题 .docx 提取流水线 v2 — 正确处理合卷和分离卷的答案格式

答案格式:
  合卷(真题及答案): "参考答案与试题解析" → 每题有 "故选X"
  分离卷: 空白卷无答案, 解析卷有 "【答案】X" 或 "故选X"

用法:
  python scripts/extract_v2.py                     # 全量提取
  python scripts/extract_v2.py --dry-run            # 测试 (每科1文件)
  python scripts/extract_v2.py --subject math       # 单学科
"""
import sys, os, re, json, argparse
from pathlib import Path
from collections import defaultdict
from docx import Document

sys.stdout.reconfigure(encoding="utf-8", errors="replace")

EXAM_DIR = Path("exam-papers")
SUBJECT_DIRS = {
    "chinese": "1.深圳中考语文（2008-2024）",
    "math": "2.深圳中考数学（2008-2025）",
    "english": "3.深圳中考英语（2008-2024）无音频",
    "physics": "4.深圳中考物理（2008-2025）少10",
    "chemistry": "5.深圳中考化学（2008-2025）少22",
    "politics": "6.深圳中考道法（2020-2023）少21年",
    "history": "7.深圳中考历史（2008-2025）少21少24",
}


def classify_paper(filename: str) -> str:
    if "空白卷" in filename:
        return "blank"
    if "解析卷" in filename or "解析" in filename:
        return "answer"
    if "答案" in filename:
        return "combined"
    return "blank"


def extract_year(filename: str) -> int | None:
    m = re.search(r'(\d{4})年', filename)
    return int(m.group(1)) if m else None


def extract_answers_from_section(paragraphs: list) -> dict:
    """从解析区提取答案映射 {题号: 答案字母}"""
    answers = {}
    current_qnum = None

    for text in paragraphs:
        # 检测题号 "1．（2分）" 或 "1．" 或 "1."
        m = re.match(r'(\d+)[\.\、．）)]', text)
        if m:
            current_qnum = int(m.group(1))

        # 检测答案 "故选B" "【答案】A" 等
        if current_qnum:
            ans_m = re.search(r'故选[：:]?\s*([A-D])', text)
            if not ans_m:
                ans_m = re.search(r'【答案】\s*([A-D])', text)
            if not ans_m:
                ans_m = re.search(r'答案[：:]\s*([A-D])', text)
            if ans_m:
                answers[current_qnum] = ans_m.group(1)

    # 如果上面没匹配到，尝试 "题号.答案字母" 格式 (如 "1.D   2.A   3.D")
    if len(answers) < 3:
        for text in paragraphs:
            # 匹配 "1.D" "2.A" "15.B" 等
            pairs = re.findall(r'(\d+)\s*[\.、]\s*([A-D])\b', text)
            for qnum_str, letter in pairs:
                qnum = int(qnum_str)
                if qnum not in answers:
                    answers[qnum] = letter

    return answers


def extract_questions_from_paragraphs(paragraphs: list, subject: str, source: str, year: int) -> list[dict]:
    """从段落列表提取选择题（题目+选项）"""
    questions = []
    current_q = None
    current_opts = []
    q_counter = 0

    for text in paragraphs:
        text = text.strip()
        if not text:
            continue

        # 跳过试卷说明
        if re.match(r'^\d+[\.\、．）)]\s*(?:答题前|全卷|作答|考试|本卷|注意|请|说明)', text):
            continue

        # 检测题号开头
        if re.match(r'^\d+[\.\、．）)]', text) and len(text) > 8:
            # 保存上一题
            if current_q and len(current_opts) >= 2:
                current_q["options"] = current_opts
                questions.append(current_q)

            q_counter += 1
            current_q = {
                "question": text,
                "subject": subject,
                "year": year,
                "source": source,
            }
            current_opts = []
            continue

        # 检测选项行
        if current_q:
            opt_m = re.match(r'^([A-D])[\.\、．）)]\s*(.+)', text)
            if opt_m:
                # 拆分合并的选项: "A．xxx B．yyy" → ["A．xxx", "B．yyy"]
                if re.search(r'\s+[B-D][\.\、．）)]', text):
                    parts = re.split(r'\s+(?=[A-D][\.\、．）)])', text)
                    current_opts.extend([p.strip() for p in parts if re.match(r'^[A-D][\.\、．）)]', p.strip())])
                else:
                    current_opts.append(text)
                continue
            # 纯选项字母
            opt_m2 = re.match(r'^([A-D])[\.\、．）)]\s*$', text)
            if opt_m2:
                current_opts.append(text)
                continue

        # 其他文本：如果还没开始收集选项，追加到题干
        if current_q and not current_opts:
            current_q["question"] += " " + text

    # 保存最后一题
    if current_q and len(current_opts) >= 2:
        current_q["options"] = current_opts
        questions.append(current_q)

    # 过滤无效
    valid = []
    for q in questions:
        opts = q.get("options", [])
        if len(opts) < 2:
            continue
        if len(q["question"]) < 8:
            continue
        # 确保选项格式正确（以A-D开头）
        clean_opts = []
        for o in opts:
            if re.match(r'^[A-D][\.\、．）)]', o):
                clean_opts.append(o)
        if len(clean_opts) < 2:
            continue
        q["options"] = clean_opts
        valid.append(q)

    return valid


def process_combined_file(filepath: Path, subject: str) -> list[dict]:
    """处理合卷格式 (题目和答案在同一文件)"""
    doc = Document(str(filepath))
    year = extract_year(filepath.name)

    # 收集所有段落文本
    all_paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    # 找到"参考答案"分隔线 (多种变体)
    answer_start = len(all_paras)
    answer_markers = ["参考答案", "试题解析", "答案与解析", "答案及解析",
                      "参考答案与试题解析", "参考答案及解析"]
    for i, text in enumerate(all_paras):
        for marker in answer_markers:
            if marker in text:
                answer_start = i
                break
        if answer_start < len(all_paras):
            break

    # 如果没找到答案分隔线，从后20%开始搜索 (答案通常在末尾)
    if answer_start >= len(all_paras):
        answer_start = max(0, int(len(all_paras) * 0.8))

    # 分离题目区和答案区
    question_paras = all_paras[:answer_start]
    answer_paras = all_paras[answer_start:]

    # 提取题目
    questions = extract_questions_from_paragraphs(question_paras, subject, filepath.name, year)

    # 提取答案
    answers = extract_answers_from_section(answer_paras)

    # 匹配答案
    for i, q in enumerate(questions):
        qnum = i + 1  # 按顺序：第N题对应解析区的第N题
        if qnum in answers:
            q["answer"] = answers[qnum]

    return questions


def process_separate_files(blank_file: Path, answer_file: Path | None, subject: str) -> list[dict]:
    """处理分离格式 (空白卷 + 解析卷)"""
    year = extract_year(blank_file.name)

    # 提取空白卷题目
    blank_doc = Document(str(blank_file))
    blank_paras = [p.text.strip() for p in blank_doc.paragraphs if p.text.strip()]
    questions = extract_questions_from_paragraphs(blank_paras, subject, blank_file.name, year)

    # 如果有解析卷，提取答案
    if answer_file and answer_file.exists():
        answer_doc = Document(str(answer_file))
        answer_paras = [p.text.strip() for p in answer_doc.paragraphs if p.text.strip()]
        answers = extract_answers_from_section(answer_paras)

        for i, q in enumerate(questions):
            qnum = i + 1
            if qnum in answers:
                q["answer"] = answers[qnum]

    return questions


def process_subject(subject: str, dry_run: bool = False) -> list[dict]:
    dir_name = SUBJECT_DIRS.get(subject)
    if not dir_name:
        return []

    subject_dir = EXAM_DIR / dir_name
    if not subject_dir.exists():
        return []

    all_docs = sorted(subject_dir.glob("*.docx"))
    if dry_run:
        all_docs = all_docs[:3]

    # 分类
    combined = [f for f in all_docs if classify_paper(f.name) == "combined"]
    blanks = [f for f in all_docs if classify_paper(f.name) == "blank"]
    answers_map = {extract_year(f.name): f for f in all_docs if classify_paper(f.name) == "answer"}

    stats = f"合卷:{len(combined)} 空白:{len(blanks)} 解析:{len(answers_map)}"
    print(f"\n{'='*60}")
    print(f"{subject}: {stats}")
    print(f"{'='*60}")

    all_qs = []

    # 处理合卷
    for f in combined:
        qs = process_combined_file(f, subject)
        with_ans = sum(1 for q in qs if "answer" in q)
        print(f"  [合卷] {f.name[:50]} → {len(qs)}题 (含答案:{with_ans})")
        all_qs.extend(qs)

    # 处理分离卷
    for blank_f in blanks:
        year = extract_year(blank_f.name)
        answer_f = answers_map.get(year)
        qs = process_separate_files(blank_f, answer_f, subject)
        with_ans = sum(1 for q in qs if "answer" in q)
        print(f"  [分离] {blank_f.name[:50]} → {len(qs)}题 (含答案:{with_ans})")
        all_qs.extend(qs)

    return all_qs


def clean_question(q: dict) -> dict:
    """清洗题目文本"""
    text = q["question"]
    text = re.sub(r'【答案】\s*[A-D]', '', text)
    text = re.sub(r'【解析】.*?(?=$)', '', text)
    text = re.sub(r'【分析】.*?(?=$)', '', text)
    text = re.sub(r'【解答】.*?(?=$)', '', text)
    text = re.sub(r'故选[：:]?\s*[A-D]', '', text)
    text = re.sub(r'[(（]\d+分[)）]', '', text)
    text = re.sub(r'[(（]\d{4}[•·.]?[^)）]*[)）]', '', text)  # 移除 "(2008•深圳)" 引用
    text = re.sub(r'\s+', ' ', text).strip()
    q["question"] = text

    # 清洗选项
    q["options"] = [re.sub(r'\s+', ' ', o).strip() for o in q.get("options", [])]

    # 确保 answer 是单个字母
    ans = q.get("answer", "")
    if ans and len(ans) > 1 and ans[0] in "ABCD":
        q["answer"] = ans[0]

    return q


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--subject", default=None)
    parser.add_argument("--dry-run", action="store_true")
    parser.add_argument("--output", default="data/extracted_questions.json")
    args = parser.parse_args()

    subjects = [args.subject] if args.subject else list(SUBJECT_DIRS.keys())
    all_results = []

    for subject in subjects:
        qs = process_subject(subject, dry_run=args.dry_run)
        for q in qs:
            clean_question(q)
        all_results.extend(qs)
        total = len(qs)
        with_ans = sum(1 for q in qs if "answer" in q)
        print(f"  → {subject}: {total}题 (含答案:{with_ans})")

    # 去重
    seen = set()
    unique = []
    for q in all_results:
        key = q["question"][:60]
        if key not in seen:
            seen.add(key)
            unique.append(q)

    # 统计
    by_subj = defaultdict(list)
    for q in unique:
        by_subj[q["subject"]].append(q)

    print(f"\n{'='*60}")
    print(f"提取完成!")
    print(f"  总提取: {len(all_results)} → 去重后: {len(unique)}")
    for subj in subjects:
        qs = by_subj.get(subj, [])
        with_ans = sum(1 for q in qs if "answer" in q)
        print(f"  {subj:12s}: {len(qs):4d}题 (含答案:{with_ans})")

    os.makedirs(os.path.dirname(args.output), exist_ok=True)
    with open(args.output, "w", encoding="utf-8") as f:
        json.dump(unique, f, ensure_ascii=False, indent=2)
    print(f"\n已保存: {args.output}")


if __name__ == "__main__":
    main()
