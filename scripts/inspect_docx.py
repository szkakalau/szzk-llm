# -*- coding: utf-8 -*-
"""检查 docx 文件结构，为提取流水线做准备"""
import sys, os, re
from pathlib import Path
sys.stdout.reconfigure(encoding="utf-8", errors="replace")

from docx import Document

# 各学科抽样一个文件
SAMPLES = {
    "chinese": "exam-papers/1.深圳中考语文（2008-2024）/2024年广东省深圳市中考语文真题（空白卷）.docx",
    "math": "exam-papers/2.深圳中考数学（2008-2025）/2024年广东省深圳市中考数学真题（空白卷）.docx",
    "english": "exam-papers/3.深圳中考英语（2008-2024）无音频/2024年广东省深圳市中考英语真题（空白卷）.docx",
    "physics": "exam-papers/4.深圳中考物理（2008-2025）少10/2024年广东省深圳市中考物理真题（空白卷）.docx",
    "chemistry": "exam-papers/5.深圳中考化学（2008-2025）少22/2024年广东省深圳市中考化学真题（空白卷）.docx",
}

for subject, path in SAMPLES.items():
    full_path = Path(path)
    if not full_path.exists():
        # 尝试找最近年份的文件
        parent = full_path.parent
        if parent.exists():
            files = sorted(parent.glob("*空白卷*.docx"))
            if files:
                full_path = files[-1]
            else:
                files = sorted(parent.glob("*.docx"))
                if files:
                    full_path = files[-1]

    if not full_path.exists():
        print(f"{subject}: 文件不存在 — {path}")
        continue

    doc = Document(str(full_path))
    paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    # 检测题目（以数字开头的段落）
    q_pattern = re.compile(r'^(\d+)[\.\、．）)]')
    questions = [p for p in paras if q_pattern.match(p)]

    # 检测选项行（以 A/B/C/D 开头）
    opt_pattern = re.compile(r'^[A-D][\.\、．）)]')

    # 统计图片
    img_count = 0
    math_count = 0
    ns = {"m": "http://schemas.openxmlformats.org/officeDocument/2006/math",
          "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
          "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
          "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
          "pic": "http://schemas.openxmlformats.org/drawingml/2006/picture"}
    for p in doc.paragraphs:
        for run in p.runs:
            # 检查数学公式
            maths = run._element.findall('.//{' + ns["m"] + '}oMath')
            math_count += len(maths)
            # 检查图片
            imgs = run._element.findall('.//{' + ns["wp"] + '}inline')
            imgs += run._element.findall('.//{' + ns["wp"] + '}anchor')
            img_count += len(imgs)
            # 也检查 drawing
            drawings = run._element.findall('.//{' + ns["wp"] + '}drawing')
            img_count += len(drawings)

    print(f"\n{'='*60}")
    print(f"学科: {subject}")
    print(f"文件: {full_path.name}")
    print(f"总段落: {len(doc.paragraphs)}, 非空段落: {len(paras)}, 表格: {len(doc.tables)}")
    print(f"检测到题目数: {len(questions)}")
    print(f"检测到图片/公式: 图片≈{img_count}, 数学公式≈{math_count}")
    print(f"前5题:")
    for q in questions[:5]:
        print(f"  {q[:120]}")
